from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import tempfile
import os
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


TEMP_IMAGES = []  # Track images to delete later


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_heading(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.color.rgb = RGBColor(31, 41, 55)
    run.font.size = Pt(14 if level == 2 else 18)


def add_badge(doc, label, value):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)

    if value >= 75:
        color = "22C55E"
    elif value >= 50:
        color = "F97316"
    else:
        color = "EF4444"

    set_cell_bg(cell, color)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{label}: {value}%")
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(12)

    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_chart_image(doc, title, data_dict):
    add_heading(doc, title, level=3)

    labels = list(data_dict.keys())
    values = list(data_dict.values())

    plt.figure(figsize=(6, 3))
    plt.bar(labels, values)
    plt.ylim(0, 100)
    plt.tight_layout()

    fd, path = tempfile.mkstemp(suffix=".png")
    os.close(fd)  # 🔑 CRITICAL FOR WINDOWS

    plt.savefig(path)
    plt.close()

    doc.add_picture(path, width=Inches(5))
    TEMP_IMAGES.append(path)  # delete later


def generate_docx_report(response, context):
    doc = Document()

    # Title
    title = doc.add_heading("Resume Skill Match Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(20)

    # Badge
    add_badge(doc, "Match Score", context["score"])

    doc.add_paragraph(f"Confidence Level: {context['confidence_level']}")

    add_heading(doc, "Detected Skills")
    add_bullets(doc, context["skills"])

    add_heading(doc, "Missing Skills")
    add_bullets(doc, context["missing_skills"])

    add_chart_image(doc, "ATS Scorecard", context["ats_scores"])
    add_chart_image(doc, "Section Score Breakdown", context["section_scores"])

    add_heading(doc, "Warnings & Insights")
    add_bullets(doc, [w["message"] for w in context["warnings"]])

    add_heading(doc, "Recruiter AI Feedback")
    for line in context["recruiter_feedback"].split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    # Save DOCX first
    doc.save(response)

    # 🔥 NOW SAFE TO DELETE TEMP FILES
    for img in TEMP_IMAGES:
        try:
            os.remove(img)
        except Exception:
            pass


def skill_gap_score(resume_text, jd_text):
    tfidf = TfidfVectorizer(stop_words="english")
    vectors = tfidf.fit_transform([resume_text, jd_text])
    similarity = cosine_similarity(vectors[0], vectors[1])[0][0]
    return round((1 - similarity) * 100, 2)




